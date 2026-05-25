from typing import List, Dict
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .models import Sentence, CorrectionIssue

class DocxAnnotator:
    def __init__(self):
        pass
    
    def create_comment(self, document: Document, sentence: Sentence, issue: CorrectionIssue):
        """
        Add a comment to the specific sentence in the DOCX.
        """
        # This is a simplified implementation
        # In a real implementation, you would find the specific sentence 
        # and add a comment or highlight it
        
        # For now, we'll just create a note at the end of the document
        document.add_paragraph(f"Correction for sentence {sentence.sentence_id}:")
        document.add_paragraph(f"  Issue: {issue.issue}")
        document.add_paragraph(f"  Fix: {issue.fix}")
        if issue.suggestions:
            document.add_paragraph(f"  Suggestions: {', '.join(issue.suggestions)}")
        document.add_paragraph("")
    
    def annotate_document(self, input_file_path: str, output_file_path: str, 
                         sentences: List[Sentence], issues: List[CorrectionIssue]):
        """
        Annotate a DOCX document with corrections.
        """
        # Load the original document
        doc = Document(input_file_path)
        
        # Add corrections as notes at the end for demonstration
        for issue in issues:
            # Find the sentence that matches this issue
            matching_sentence = next((s for s in sentences if s.sentence_id == issue.sentence_id), None)
            if matching_sentence:
                self.create_comment(doc, matching_sentence, issue)
        
        # Save the annotated document
        doc.save(output_file_path)